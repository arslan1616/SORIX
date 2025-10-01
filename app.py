import os
import shutil
import uuid
import json
from flask import Flask, render_template, request, jsonify, send_from_directory
import google.generativeai as genai
from dotenv import load_dotenv
from datetime import datetime
from docx import Document
from docx.shared import Inches
import cairosvg
import tempfile  # GÜVENLİ GEÇİCİ DOSYA İŞLEMLERİ İÇİN EKLENDİ

# --- Kurulum ---
load_dotenv()
app = Flask(__name__)
GENERATED_FILES_DIR = 'generated_files'
HISTORY_FILE = 'history.json'
app.config['GENERATED_FILES_DIR'] = os.path.join(os.getcwd(), GENERATED_FILES_DIR)
if not os.path.exists(GENERATED_FILES_DIR): os.makedirs(GENERATED_FILES_DIR)

# --- API Yapılandırması ---
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key: raise ValueError("GOOGLE_API_KEY bulunamadı.")
genai.configure(api_key=api_key)


# --- Yardımcı Fonksiyonlar ---
def load_history():
    if not os.path.exists(HISTORY_FILE) or os.path.getsize(HISTORY_FILE) == 0: return []
    try:
        with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except json.JSONDecodeError:
        return []


def save_history(data):
    with open(HISTORY_FILE, 'w', encoding='utf-8') as f: json.dump(data, f, ensure_ascii=False, indent=4)


# --- GÜNCELLENDİ: Çökmeyi Engelleyen create_documents Fonksiyonu ---
def create_documents(questions_data, file_id, topic_details):
    q_doc, s_doc = Document(), Document()
    q_doc.add_heading(f"{topic_details} - Sorular", 0)
    s_doc.add_heading(f"{topic_details} - Çözümler", 0)
    generated_svgs = []
    alphabet = "ABCDE"

    for i, item in enumerate(questions_data, 1):
        q_doc.add_heading(f"Soru {i}", level=1);
        q_doc.add_paragraph(item['question'])
        s_doc.add_heading(f"Soru {i}", level=1);
        s_doc.add_paragraph(item['question'])

        if item.get('svg_image') and item['svg_image'].strip():
            # tempfile ile güvenli geçici dosya oluşturma
            with tempfile.NamedTemporaryFile(mode='w+', suffix='.svg', delete=False, encoding='utf-8') as temp_svg_file:
                temp_svg_file.write(item['svg_image'])
                temp_svg_path = temp_svg_file.name

            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_png_file:
                temp_png_path = temp_png_file.name

            try:
                cairosvg.svg2png(url=temp_svg_path, write_to=temp_png_path)
                q_doc.add_picture(temp_png_path, width=Inches(4.0))
                s_doc.add_picture(temp_png_path, width=Inches(4.0))

                # Kalıcı SVG dosyasını kaydet
                downloadable_svg_filename = f'gorsel_{file_id}_{i}.svg'
                permanent_svg_path = os.path.join(app.config['GENERATED_FILES_DIR'], downloadable_svg_filename)
                shutil.move(temp_svg_path, permanent_svg_path)
                generated_svgs.append(downloadable_svg_filename)

            except Exception as e:
                print(f"SVG -> PNG dönüşüm hatası: {e}")

            finally:
                # İşlem bittikten sonra geçici dosyaları temizle
                if os.path.exists(temp_svg_path): os.remove(temp_svg_path)
                if os.path.exists(temp_png_path): os.remove(temp_png_path)

        # ... (Şıkları ve çözümleri ekleme kısmı aynı) ...
        if item.get('options'):
            for idx, option in enumerate(item['options']): q_doc.add_paragraph(
                f"{alphabet[idx]}) {option}"); s_doc.add_paragraph(f"{alphabet[idx]}) {option}")
        q_doc.add_paragraph()
        s_doc.add_heading("Çözüm Adımları:", level=2)
        for step in item['solution_steps']: s_doc.add_paragraph(step, style='List Bullet')
        correct_answer_text = item['options'][item['correct_answer_index']] if item.get('options') else item.get(
            'answer', '')
        s_doc.add_paragraph(f"\nDoğru Cevap: {correct_answer_text}")
        s_doc.add_page_break()

    q_word_path = f'sorular_{file_id}.docx';
    s_word_path = f'cozumler_{file_id}.docx'
    q_doc.save(os.path.join(app.config['GENERATED_FILES_DIR'], q_word_path))
    s_doc.save(os.path.join(app.config['GENERATED_FILES_DIR'], s_word_path))
    return {"questions_word": q_word_path, "solutions_word": s_word_path, "svg_files": generated_svgs}


# --- Yapay Zeka Fonksiyonu (Değişiklik yok) ---
def generate_ai_questions_in_batch(topic, sub_topic, count, is_visual, question_type):
    model = genai.GenerativeModel('gemini-pro-latest')

    if question_type == 'multiple_choice':
        type_instruction = """
            İstenen çıktı formatı:
            - "question": Soru metni.
            - "options": 4 adet metin içeren bir dizi (array). Şıklardan sadece biri doğru olmalı, diğerleri mantıklı çeldiriciler olmalıdır.
            - "correct_answer_index": Doğru şıkkın dizindeki sırası (0, 1, 2 veya 3).
            - "solution_steps": Doğru cevabın neden doğru olduğunu açıklayan adımlar.
            - "svg_image": Varsa SVG metni, yoksa boş string.
            """
    else:  # 'classic'
        type_instruction = """
            İstenen çıktı formatı:
            - "question": Soru metni.
            - "answer": Sorunun kısa ve net cevabı.
            - "solution_steps": Cevaba nasıl ulaşıldığını açıklayan adımlar.
            - "svg_image": Varsa SVG metni, yoksa boş string.
            """

    visual_instruction = "Eğer soru için bir görsel uygunsa, 'svg_image' alanını geçerli SVG koduyla doldur. Uygun değilse boş bırak." if is_visual else ""

    prompt = f"""
        Sen, Türkiye'deki eğitim müfredatına uygun, lise seviyesinde sorular hazırlayan bir yapay zeka asistanısın.
        Görevin, aşağıda belirtilen konuda, istenen sayıda ve tipte soru üretmektir.
        Lütfen çıktıyı SADECE ve SADECE geçerli bir JSON formatında, "questions" adında tek bir anahtara sahip bir nesne olarak ver. Bu anahtarın değeri, istenen formatta soru nesneleri içeren bir dizi (array) olmalıdır.

        {type_instruction}
        {visual_instruction}

        Ana Konu: {topic}
        Alt Konu: {sub_topic}
        İstenen Soru Sayısı: {count}
        """
    try:
        response = model.generate_content(prompt)

        # --- EN ÖNEMLİ DÜZELTME BURADA ---
        # API'den gelen cevabın içeriğini kontrol et
        if not response.parts:
            print("API Hatası: Gemini API'sinden boş yanıt alındı (Safety filter olabilir).")
            return None

        cleaned_response_text = response.text.strip().replace('```json', '').replace('```', '')

        # JSON'a çevirmeden önce metnin boş olup olmadığını tekrar kontrol et
        if not cleaned_response_text:
            print("API Hatası: Temizlenmiş yanıt metni boş.")
            return None

        ai_data = json.loads(cleaned_response_text)
        return ai_data.get("questions", [])
    except Exception as e:
        print(f"API Hatası veya JSON Parse Hatası: {e}")
        return None


# --- Rotalar ---
@app.route('/')
def index(): return render_template('index.html')


@app.route('/history', methods=['GET'])
def get_history(): return jsonify(sorted(load_history(), key=lambda x: x['timestamp'], reverse=True))


@app.route('/generate-documents', methods=['POST'])
def generate_documents_route():
    # ... (Bu rota önceki versiyonla aynı, değişiklik yok)
    data = request.json;
    topic, sub_topic, count, is_visual, question_type = data.get('topic'), data.get('sub_topic'), data.get(
        'count'), data.get('is_visual'), data.get('question_type', 'classic')
    questions_list = generate_ai_questions_in_batch(topic, sub_topic, int(count), is_visual, question_type)
    if questions_list is None: return jsonify({"error": "Yapay zekadan soru üretilemedi."}), 500
    unique_id = uuid.uuid4().hex[:8];
    topic_details = f"{topic} - {sub_topic}";
    file_names = create_documents(questions_list, unique_id, topic_details)
    new_history_entry = {"id": unique_id, "topic": topic_details, "count": len(questions_list), "files": file_names,
                         "timestamp": datetime.now().isoformat(), "questions": questions_list}
    history = load_history();
    history.append(new_history_entry);
    save_history(history)
    return jsonify(new_history_entry)


# YENİ: Silme İşlemi İçin Rota
@app.route('/delete-item/<item_id>', methods=['DELETE'])
def delete_item(item_id):
    history = load_history()
    item_to_delete = None
    for item in history:
        if item.get('id') == item_id:
            item_to_delete = item
            break

    if item_to_delete:
        # İlişkili dosyaları sil
        for file_key in item_to_delete.get('files', {}):
            filename = item_to_delete['files'][file_key]
            if isinstance(filename, str):  # SVG listesi değilse
                filepath = os.path.join(app.config['GENERATED_FILES_DIR'], filename)
                if os.path.exists(filepath): os.remove(filepath)
            elif isinstance(filename, list):  # SVG listesi ise
                for svg_file in filename:
                    filepath = os.path.join(app.config['GENERATED_FILES_DIR'], svg_file)
                    if os.path.exists(filepath): os.remove(filepath)

        # history.json'dan kaydı sil
        updated_history = [item for item in history if item.get('id') != item_id]
        save_history(updated_history)
        return jsonify({"success": True, "message": "Kayıt ve ilişkili dosyalar silindi."})

    return jsonify({"success": False, "message": "Kayıt bulunamadı."}), 404


@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(app.config['GENERATED_FILES_DIR'], filename, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)